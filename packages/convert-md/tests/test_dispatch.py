"""Opinionated conversion dispatch — mechanism (ADR 0021 / R142)."""

from __future__ import annotations

from pathlib import Path

import docx as docx_writer
from convert_md import (
    Html2TextEngine,
    MammothEngine,
    MarkitdownEngine,
    OpenpyxlEngine,
    PymupdfLlmEngine,
    TrafilaturaEngine,
    convert,
    fallback_chain_for,
    grade,
    select_engine,
)
from openpyxl import Workbook


def test_dispatch_table_primary_engines() -> None:
    assert select_engine(Path("x.pdf")) is PymupdfLlmEngine
    assert select_engine(Path("x.docx")) is MammothEngine
    assert select_engine(Path("x.pptx")) is MarkitdownEngine
    assert select_engine(Path("x.xlsx")) is MarkitdownEngine
    assert select_engine(Path("x.html")) is TrafilaturaEngine


def test_fallback_chains() -> None:
    assert fallback_chain_for(Path("x.pdf")) == [PymupdfLlmEngine]
    assert fallback_chain_for(Path("x.docx")) == [MammothEngine, MarkitdownEngine]
    assert fallback_chain_for(Path("x.xlsx")) == [MarkitdownEngine, OpenpyxlEngine]
    assert fallback_chain_for(Path("x.html")) == [TrafilaturaEngine, Html2TextEngine]


def test_unsupported_format_returns_failed() -> None:
    result = convert(Path("/nonexistent/archive.zip"))
    assert result.fidelity == "failed"
    assert any("unsupported_format" in w for w in result.warnings)


def test_docx_round_trips_via_mammoth(tmp_path: Path) -> None:
    doc = docx_writer.Document()
    doc.add_heading("Title", level=1)
    p = doc.add_paragraph("Body paragraph with ")
    p.add_run("bold").bold = True
    p.add_run(".")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "a", "b"
    table.rows[1].cells[0].text, table.rows[1].cells[1].text = "1", "2"
    docx = tmp_path / "doc.docx"
    doc.save(str(docx))

    result = convert(docx)
    assert result.engine.startswith("mammoth@")  # mammoth front-end, then the render engine
    assert result.fidelity == "high"
    assert "Title" in result.body_markdown
    assert "**bold**" in result.body_markdown  # emphasis preserved
    assert "|" in result.body_markdown  # table preserved


def test_xlsx_converts_with_sheet_content(tmp_path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Numbers"
    ws.append(["name", "value"])
    ws.append(["alpha", 42])
    xlsx = tmp_path / "sheet.xlsx"
    wb.save(xlsx)

    result = convert(xlsx)
    assert result.engine.startswith(("markitdown@", "openpyxl@"))
    assert "alpha" in result.body_markdown or "42" in result.body_markdown


def test_html_converts_via_trafilatura_or_fallback(tmp_path: Path) -> None:
    html = tmp_path / "page.html"
    html.write_text(
        "<html><body><article><h1>Heading</h1>"
        "<p>" + ("This is a substantial paragraph of article content. " * 8) + "</p>" + "</article></body></html>",
        encoding="utf-8",
    )
    result = convert(html)
    assert result.engine.startswith(("trafilatura@", "html2text@"))
    assert result.body_markdown.strip()


def test_corrupt_docx_walks_to_fallback_engine(tmp_path: Path) -> None:
    # mammoth rejects a non-zip .docx; the dispatcher must walk to the markitdown
    # fallback rather than failing outright (chain-walk-on-failure).
    bad = tmp_path / "broken.docx"
    bad.write_bytes(b"this is not a real docx zip but has readable text")
    result = convert(bad)
    assert result.engine.startswith("markitdown@")  # primary (mammoth) was skipped


def test_fidelity_grade_failed_on_empty() -> None:
    fidelity, lost, _warnings = grade(markdown="   \n", source_size=5000)
    assert fidelity == "failed"
    assert "all" in lost


def test_fidelity_grade_partial_on_garbage() -> None:
    garbled = "a" + "�" * 50
    fidelity, _lost, warnings = grade(markdown=garbled, source_size=200)
    assert fidelity == "partial"
    assert any("garbage" in w for w in warnings)


def test_fidelity_grade_partial_on_cid_glyph_leak() -> None:
    # A raw PDF font's ToUnicode mapping failing to resolve leaks literal
    # (cid:N) tokens into the text — calibrated 2026-07-09 against real
    # markitdown/CJK bench output (bench/results/2026-07-09-findings.md).
    lines = [f"real text line {i}" if i % 3 else f"(cid:{1000 + i})" for i in range(30)]
    fidelity, lost, warnings = grade(markdown="\n".join(lines), source_size=2000)
    assert fidelity == "partial"
    assert "encoding" in lost
    assert any("cid_glyph_leak" in w for w in warnings)


def test_fidelity_grade_partial_on_shattered_table() -> None:
    # An engine that drops pipe/row syntax entirely and emits one cell per
    # line — calibrated against unstructured's real shattered-table output
    # (0.562 numeric-line fraction vs. 0.260 for the highest legitimate case).
    lines = [str(1000 + i) if i % 2 else f"label {i}" for i in range(30)]
    fidelity, lost, warnings = grade(markdown="\n".join(lines), source_size=2000)
    assert fidelity == "partial"
    assert "table_structure" in lost
    assert any("shattered_table" in w for w in warnings)


def test_fidelity_grade_high_on_real_markdown_table_despite_numeric_density() -> None:
    # A genuine markdown table is numeric-dense by nature but every cell sits
    # inside `| … |` syntax — must not trip the shattered-table check.
    rows = "\n".join(f"| {i} | {i * 2} | {i * 3} |" for i in range(20))
    fidelity, _lost, _warnings = grade(markdown=f"| a | b | c |\n|---|---|---|\n{rows}\n", source_size=2000)
    assert fidelity == "high"


def test_fidelity_grade_high_on_clean_text() -> None:
    fidelity, _lost, _warnings = grade(markdown="# Clean\n\nA fine document.\n", source_size=120)
    assert fidelity == "high"


def test_pdf_routes_to_pymupdf4llm_only() -> None:
    assert select_engine(Path("paper.pdf")) is PymupdfLlmEngine
    assert fallback_chain_for(Path("paper.pdf")) == [PymupdfLlmEngine]


def _minimal_pdf() -> bytes:
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length 70 >>\nstream\nBT /F1 24 Tf 72 700 Td (Hello docling two) Tj ET\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = b"%PDF-1.4\n"
    offsets: list[int] = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objs) + 1}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
    return out


def test_pdf_round_trip_via_pymupdf4llm(tmp_path: Path) -> None:
    """pymupdf4llm (the only `.pdf` engine as of v0.5.0) converts a real PDF through the real dispatch chain."""
    pdf = tmp_path / "hello.pdf"
    pdf.write_bytes(_minimal_pdf())
    result = convert(pdf)
    assert result.engine.startswith("pymupdf4llm@")
    assert result.fidelity == "high"
    assert "Hello docling two" in result.body_markdown
